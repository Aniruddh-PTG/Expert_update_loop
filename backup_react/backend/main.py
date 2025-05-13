import os
import base64
import requests
from flask import Flask, request, jsonify
from dotenv import load_dotenv
import io
from docx import Document
from flask_cors import CORS
import boto3
import json

from flask_cors import CORS
app = Flask(__name__)
CORS(app)

load_dotenv()
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
REPO = os.getenv("REPO")
BRANCH = os.getenv("BRANCH", "main")
AWS_REGION = os.getenv("AWS_DEFAULT_REGION") or os.getenv("AWS_REGION")

app = Flask(__name__)
CORS(app)

def get_file_from_github(file_path):
    api_url = f"https://api.github.com/repos/{REPO}/contents/{file_path}?ref={BRANCH}"
    r = requests.get(api_url, headers={"Authorization": f"token {GITHUB_TOKEN}"})
    if r.status_code == 200:
        content = base64.b64decode(r.json()["content"])
        return content
    else:
        return None

def github_upload(file_path, file_content, commit_message):
    api_url = f"https://api.github.com/repos/{REPO}/contents/{file_path}"
    r = requests.get(api_url + f"?ref={BRANCH}", headers={"Authorization": f"token {GITHUB_TOKEN}"})
    sha = r.json().get("sha") if r.status_code == 200 else None
    data = {
        "message": commit_message,
        "content": base64.b64encode(file_content).decode(),
        "branch": BRANCH
    }
    if sha:
        data["sha"] = sha
    resp = requests.put(api_url, json=data, headers={"Authorization": f"token {GITHUB_TOKEN}"})
    return resp.status_code in (200, 201), resp.json()

def extract_text_from_docx(content_bytes):
    doc = Document(io.BytesIO(content_bytes))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def get_claude_summary(old_text, new_text):
    prompt = (
        "You are an expert document reviewer. Summarize the main changes between the following two versions of a document. "
        "Focus only on what changed, in plain English, in 1-3 sentences."
        "\n\nHuman: Here are the two versions:\n"
        "---\n"
        "ORIGINAL:\n"
        f"{old_text}\n"
        "---\n"
        "MODIFIED:\n"
        f"{new_text}\n"
        "---\n"
        "Please summarize the changes."
        "\n\nAssistant:"
    )
    bedrock = boto3.client("bedrock-runtime", region_name=AWS_REGION)
    body = json.dumps({
        "prompt": prompt,
        "max_tokens_to_sample": 256,
        "temperature": 0.2,
    }).encode("utf-8")
    response = bedrock.invoke_model(
        modelId="anthropic.claude-v2",
        body=body,
        accept="application/json",
        contentType="application/json"
    )
    result = response["body"].read().decode()
    # Extract only the 'completion' field from the JSON response
    try:
        import re
        import json as pyjson
        result_json = pyjson.loads(result)
        summary = result_json.get("completion", "").strip()
    except Exception:
        summary = result.strip()
    return summary

@app.route("/upload", methods=["POST"])
def upload():
    file = request.files["file"]
    orig_name = file.filename
    orig_bytes = file.read()
    ok, resp = github_upload(orig_name, orig_bytes, f"Upload {orig_name} from React")
    return jsonify({"success": ok, "response": resp})

@app.route("/sync-modified", methods=["POST"])
def sync_modified():
    orig_name = "Password Test Case Updated.docx"
    mod_name = "Password Test Case Updated modified.docx"
    content = get_file_from_github(orig_name)
    if content:
        # Extract text for summary
        orig_text = extract_text_from_docx(content)
        # Prepare a doc with the modification for summary generation
        doc_for_summary = Document(io.BytesIO(content))
        doc_for_summary.add_paragraph("MODIFIED BY BACKEND")
        out_buf = io.BytesIO()
        doc_for_summary.save(out_buf)
        out_buf.seek(0)
        mod_content_for_summary = out_buf.read()
        mod_text = extract_text_from_docx(mod_content_for_summary)
        summary = ""
        try:
            summary = get_claude_summary(orig_text, mod_text)
        except Exception as e:
            summary = f"[LLM error: {e}]"
        # Now, create the final modified docx with summary at the top
        doc = Document(io.BytesIO(content))
        # Insert summary as first paragraph
        p = doc.add_paragraph()
        run = p.add_run(summary)
        try:
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(0, 176, 80)  # Green
        except Exception:
            pass
        # Move the summary paragraph to the top
        doc._body._element.insert(0, p._element)
        # Add the modification at the end
        doc.add_paragraph("MODIFIED BY BACKEND")
        out_buf = io.BytesIO()
        doc.save(out_buf)
        out_buf.seek(0)
        mod_content = out_buf.read()
        ok, resp = github_upload(mod_name, mod_content, f"Update {mod_name} from React (with modification and summary)")
        return jsonify({"success": ok, "response": resp, "summary": summary})
    else:
        return jsonify({"success": False, "error": f"Could not fetch {orig_name} from GitHub."})

@app.route("/summarize-docx", methods=["POST"])
def summarize_docx():
    if "original" not in request.files or "modified" not in request.files:
        return jsonify({"success": False, "error": "Both 'original' and 'modified' files are required."}), 400
    orig_file = request.files["original"]
    mod_file = request.files["modified"]
    orig_bytes = orig_file.read()
    mod_bytes = mod_file.read()
    orig_text = extract_text_from_docx(orig_bytes)
    mod_text = extract_text_from_docx(mod_bytes)
    summary = ""
    try:
        summary = get_claude_summary(orig_text, mod_text)
    except Exception as e:
        summary = f"[LLM error: {e}]"
    return jsonify({"success": True, "summary": summary})

if __name__ == "__main__":
    app.run(debug=True, port=5000) 