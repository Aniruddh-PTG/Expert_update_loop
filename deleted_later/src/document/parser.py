from docx import Document
from typing import List, Dict
import io

class DocumentParser:
    @staticmethod
    def parse_docx(content: bytes) -> str:
        """
        Parse a .docx file content into plain text
        """
        try:
            doc = Document(io.BytesIO(content))
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            raise Exception(f"Error parsing .docx file: {str(e)}")
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """
        Extract sections from the document text based on headings
        """
        sections = {}
        current_section = "Introduction"
        current_content = []
        
        for line in text.split("\n"):
            # Simple heading detection (can be improved)
            if line.strip().isupper() or line.strip().endswith(":"):
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = line.strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Add the last section
        if current_content:
            sections[current_section] = "\n".join(current_content)
            
        return sections 